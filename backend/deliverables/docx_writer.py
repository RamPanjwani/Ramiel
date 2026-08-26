"""Word Document (.docx) Deliverable Generator for Ramiel.

Phase 5: Deliverable Generation.
Generates structured Word documents (such as approval notes, technical inspection summaries,
and executive briefs) from agent findings using python-docx and local template files.
"""

from __future__ import annotations

import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor


class DocxWriter:
    """Generates formatted Microsoft Word (.docx) documents."""

    def __init__(self, default_output_dir: str | Path = "data/uploads") -> None:
        self.output_dir = Path(default_output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        findings: dict[str, Any],
        filename: str | None = None,
        template: str | None = None,
    ) -> str:
        """Generate an executive approval note or report in Word format.

        Args:
            findings: Dictionary with:
                - 'title': Document title (str)
                - 'subject': Document subject / memo header (str)
                - 'executive_summary': Executive summary text (str)
                - 'sections': List of dicts with 'heading' and 'content' (str)
                - 'table': Optional dict with 'headers' (list[str]) and 'rows' (list[list[Any]])
                - 'sign_off': Optional list of sign-off roles / names
            filename: Custom filename (defaults to generated UUID filename)
            template: Optional path to template .docx

        Returns:
            The filesystem path to the saved .docx file.
        """
        if template:
            template_path = Path(template)
            if not template_path.exists():
                raise FileNotFoundError(f"Template not found: {template_path}")
            doc = Document(str(template_path))
        else:
            doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x16, 0x1B, 0x1E)

        # Title
        title_text = findings.get("title", "EXECUTIVE APPROVAL NOTE")
        title_para = doc.add_heading(title_text, level=0)
        title_para.alignment = 0  # Left aligned

        # Metadata Header Block (Memo format)
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.autofit = False
        meta_table.columns[0].width = Inches(1.5)
        meta_table.columns[1].width = Inches(5.0)

        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        headers = [
            ("DATE:", date_str),
            ("TO:", findings.get("to", "Executive Leadership / Operations Committee")),
            ("FROM:", findings.get("from", "Sovereign Engineering Review System")),
            ("SUBJECT:", findings.get("subject", "Technical Review & Recommendation")),
        ]
        for idx, (label, val) in enumerate(headers):
            row = meta_table.rows[idx]
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(label)
            r0.bold = True
            r0.font.size = Pt(10)
            p1 = row.cells[1].paragraphs[0]
            r1 = p1.add_run(val)
            r1.font.size = Pt(10)

        doc.add_paragraph()  # Spacer

        # Executive Summary
        if "executive_summary" in findings:
            doc.add_heading("1. Executive Summary", level=1)
            doc.add_paragraph(findings["executive_summary"])

        # Sections
        sections = findings.get("sections", [])
        for idx, sec in enumerate(sections, start=2):
            heading = sec.get("heading", f"Section {idx}")
            content = sec.get("content", "")
            doc.add_heading(f"{idx}. {heading}", level=1)
            doc.add_paragraph(content)

        # Table data if provided
        table_data = findings.get("table")
        if table_data and "headers" in table_data and "rows" in table_data:
            doc.add_heading("Key Findings & Inspection Data", level=2)
            headers_list = table_data["headers"]
            rows_list = table_data["rows"]

            t = doc.add_table(rows=len(rows_list) + 1, cols=len(headers_list))
            t.style = "Table Grid"

            # Headers
            for col_idx, header in enumerate(headers_list):
                cell = t.rows[0].cells[col_idx]
                p = cell.paragraphs[0]
                run = p.add_run(str(header))
                run.bold = True
                run.font.size = Pt(10)

            # Rows
            for row_idx, row_values in enumerate(rows_list):
                for col_idx, val in enumerate(row_values):
                    cell = t.rows[row_idx + 1].cells[col_idx]
                    p = cell.paragraphs[0]
                    p.add_run(str(val) if val is not None else "").font.size = Pt(9.5)

            doc.add_paragraph()

        # Sign-off / Approval Block
        sign_offs = findings.get(
            "sign_off",
            ["Reviewed By: Senior Lead Engineer", "Approved By: Plant General Manager"],
        )
        if sign_offs:
            doc.add_heading("Sign-off & Approvals", level=2)
            for signee in sign_offs:
                p = doc.add_paragraph()
                p.add_run(
                    f"☐  {signee}\n    Date: _______________    Signature: ______________________\n"
                )

        # Save document
        out_filename = filename or f"approval_note_{uuid.uuid4().hex[:6]}.docx"
        out_path = self.output_dir / out_filename
        doc.save(str(out_path))
        return str(out_path)
