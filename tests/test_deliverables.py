"""Deliverables Engine Tests — Phase 5.

Validates:
1. DocxWriter Word document generation with memos, tables, and sign-offs.
2. PptxWriter PowerPoint presentation generation with slide layouts.
3. XlsxWriter Excel workbook generation with multi-sheet styled tables.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import docx
import openpyxl
import pptx

from backend.deliverables.docx_writer import DocxWriter
from backend.deliverables.pptx_writer import PptxWriter
from backend.deliverables.xlsx_writer import XlsxWriter


class TestDocxWriter:
    """Test Word document generation."""

    def test_generate_approval_note(self) -> None:
        out_dir = Path(tempfile.mkdtemp())
        writer = DocxWriter(default_output_dir=out_dir)

        findings = {
            "title": "PLANT OVERHAUL APPROVAL NOTE",
            "subject": "Turbine Valve Replacement Recommendation",
            "executive_summary": "Inspection revealed significant wear on Valve V-102.",
            "sections": [
                {
                    "heading": "Technical Assessment",
                    "content": "Pressure drop exceeded tolerance limits by 14.5%.",
                },
            ],
            "table": {
                "headers": ["Tag", "Parameter", "Limit", "Observed", "Status"],
                "rows": [
                    ["V-101", "Pressure", "15 bar", "12.5 bar", "OK"],
                    ["V-102", "Pressure", "15 bar", "17.2 bar", "FAIL"],
                ],
            },
        }

        path_str = writer.generate(findings, filename="test_approval.docx")
        assert Path(path_str).exists()

        # Validate docx structure
        doc = docx.Document(path_str)
        headings = [p.text for p in doc.paragraphs if p.text]
        assert "PLANT OVERHAUL APPROVAL NOTE" in headings
        assert len(doc.tables) >= 2


class TestPptxWriter:
    """Test PowerPoint presentation deck generation."""

    def test_generate_presentation(self) -> None:
        out_dir = Path(tempfile.mkdtemp())
        writer = PptxWriter(default_output_dir=out_dir)

        content = {
            "title": "Quarterly Operations Review",
            "subtitle": "Turbine & Valve Analysis",
            "slides": [
                {
                    "title": "Key Observations",
                    "bullet_points": [
                        "Turbine efficiency maintained at 94.2%",
                        "Valve V-102 scheduled for maintenance",
                        "Zero unplanned outages in Q3",
                    ],
                },
            ],
        }

        path_str = writer.generate(content, filename="test_deck.pptx")
        assert Path(path_str).exists()

        # Validate pptx structure
        prs = pptx.Presentation(path_str)
        assert len(prs.slides) == 2


class TestXlsxWriter:
    """Test Excel calculation workbook generation."""

    def test_generate_workbook(self) -> None:
        out_dir = Path(tempfile.mkdtemp())
        writer = XlsxWriter(default_output_dir=out_dir)

        data = {
            "sheets": [
                {
                    "sheet_name": "Hydraulic_Calc",
                    "headers": ["Pipe_ID", "Length_m", "Diameter_mm", "Head_Loss_m"],
                    "rows": [
                        ["P-001", 120.0, 150.0, 2.45],
                        ["P-002", 85.0, 100.0, 3.12],
                    ],
                }
            ]
        }

        path_str = writer.generate(data, filename="test_calc.xlsx")
        assert Path(path_str).exists()

        # Validate openpyxl structure
        wb = openpyxl.load_workbook(path_str)
        assert "Hydraulic_Calc" in wb.sheetnames
        ws = wb["Hydraulic_Calc"]
        assert ws.cell(row=1, column=1).value == "Pipe_ID"
        assert ws.cell(row=2, column=4).value == 2.45
